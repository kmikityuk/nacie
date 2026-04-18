#!/usr/bin/env python3
from __future__ import annotations

import re
import shutil
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_BREAK
from docx.text.paragraph import Paragraph


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

SCRIPT_PATH = Path(__file__).resolve()
ROOT = SCRIPT_PATH.parents[2]

TEMPLATE_PATH = ROOT / "00-admin" / "template" / "tecdoc-template.docx"
SOURCE_ROOT = ROOT / "01-source"
OUTPUT_PATH = ROOT / "06-master" / "01-nacie_tecdoc_master_skeleton.docx"


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class SourceItem:
    """
    One source document discovered under 01-source.
    """
    section: str
    item_no: Optional[int]
    short_name: str
    version: Optional[str]
    path: Path


# ---------------------------------------------------------------------------
# Low-level DOCX helpers
# ---------------------------------------------------------------------------

def clear_document_body_preserve_sections(doc: Document) -> None:
    """
    Remove all body content but keep the final section properties node.

    This preserves template styles, page setup, headers/footers, etc.
    """
    body = doc.element.body
    children = list(body.iterchildren())

    for child in children:
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def add_page_break(doc: Document) -> Paragraph:
    """
    Insert an explicit page break in its own paragraph.
    """
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    return p


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def find_first_existing_style(
    doc: Document,
    candidates: list[str],
    fallback: str = "Normal",
) -> str:
    """
    Return the first style name that exists in the document.
    """
    available = {style.name for style in doc.styles}
    for name in candidates:
        if name in available:
            return name
    return fallback


def add_styled_paragraph(doc: Document, text: str, style_name: str) -> Paragraph:
    """
    Add a paragraph with the requested style.
    """
    p = doc.add_paragraph(text)
    p.style = style_name
    return p


# ---------------------------------------------------------------------------
# Source scanning helpers
# ---------------------------------------------------------------------------

def parse_source_filename(path: Path) -> tuple[Optional[int], str, Optional[str]]:
    """
    Parse filenames like:
      01-chapter01-v01.docx
      10-conclusions-v00.docx
      01-ANL-v01.docx
      annexI-technical-specification-v00.docx
    """
    name = path.name

    pattern_with_item = re.compile(
        r"^(?P<item>\d{2})-(?P<short>.+)-v(?P<ver>\d{2})\.(docx|pdf|txt|md)$",
        re.IGNORECASE,
    )
    pattern_without_item = re.compile(
        r"^(?P<short>.+)-v(?P<ver>\d{2})\.(docx|pdf|txt|md)$",
        re.IGNORECASE,
    )

    match = pattern_with_item.fullmatch(name)
    if match:
        return int(match.group("item")), match.group("short"), match.group("ver")

    match = pattern_without_item.fullmatch(name)
    if match:
        return None, match.group("short"), match.group("ver")

    return None, path.stem, None


def titleize_slug(slug: str) -> str:
    """
    Convert filename slug into a readable heading.
    """
    special = {
        "chapter01": "Chapter 1",
        "chapter02": "Chapter 2",
        "chapter03": "Chapter 3",
        "chapter04": "Chapter 4",
        "chapter05": "Chapter 5",
        "chapter06": "Chapter 6",
        "chapter07": "Chapter 7",
        "chapter08": "Chapter 8",
        "chapter09": "Chapter 9",
        "conclusions": "Conclusions",
        "references": "References",
        "abbreviations": "List of Abbreviations",
        "annexI-technical-specification": "Annex I. Technical Specification",
    }
    if slug in special:
        return special[slug]

    return slug.replace("-", " ").replace("_", " ")


def scan_source_items(section_dir: Path, section_name: str) -> list[SourceItem]:
    """
    Scan one source section directory and return discovered files.
    """
    if not section_dir.exists():
        return []

    items: list[SourceItem] = []
    for path in sorted(section_dir.iterdir()):
        if not path.is_file():
            continue
        if path.suffix.lower() != ".docx":
            continue

        item_no, short_name, version = parse_source_filename(path)
        items.append(
            SourceItem(
                section=section_name,
                item_no=item_no,
                short_name=short_name,
                version=version,
                path=path,
            )
        )

    items.sort(
        key=lambda x: (
            x.item_no is None,
            x.item_no if x.item_no is not None else 9999,
            x.short_name,
        )
    )
    return items


# ---------------------------------------------------------------------------
# Skeleton writing
# ---------------------------------------------------------------------------

def add_blank_body_placeholder(doc: Document, normal_style: str) -> None:
    """
    Add one visible placeholder body paragraph.
    """
    add_styled_paragraph(doc, "Body text placeholder.", normal_style)


def add_title_page(
    doc: Document,
    tecdoc_style: str,
    title_style: str,
    subtitle_style: str,
) -> None:
    """
    Add the title page content.

    This deliberately does not use H1. The title page is front matter and should
    use dedicated front-page styles where available.
    """
    add_styled_paragraph(doc, "IAEA-TECDOC-XXXX", tecdoc_style)
    add_styled_paragraph(doc, "TITLE IS HERE", title_style)
    add_styled_paragraph(doc, "subtitle", subtitle_style)


def add_foreword_page(doc: Document, h1_style: str, normal_style: str) -> None:
    """
    Add the Foreword on its own page.
    """
    add_styled_paragraph(doc, "Foreword", h1_style)
    add_blank_body_placeholder(doc, normal_style)


def add_toc_page(doc: Document, h1_style: str, normal_style: str) -> None:
    """
    Add a safe TOC placeholder page.

    A real Word TOC field can be inserted later once the skeleton is stable.
    """
    add_styled_paragraph(doc, "CONTENTS", h1_style)
    add_styled_paragraph(doc, "[Insert and update Table of Contents in Word]", normal_style)


def add_chapter_skeleton(
    doc: Document,
    chapters: list[SourceItem],
    h1_style: str,
    normal_style: str,
) -> None:
    """
    Add headings for chapter files in source order, each on a new page.
    """
    for index, item in enumerate(chapters):
        if index > 0:
            add_page_break(doc)

        heading = titleize_slug(item.short_name)
        add_styled_paragraph(doc, heading, h1_style)
        add_blank_body_placeholder(doc, normal_style)


def add_annex_group(
    doc: Document,
    group_title: str,
    items: list[SourceItem],
    h1_style: str,
    h2_style: str,
    normal_style: str,
) -> None:
    """
    Add one annex group heading and child headings for its source files.

    Each top-level annex group starts on a new page.
    """
    if not items:
        return

    add_page_break(doc)
    add_styled_paragraph(doc, group_title, h1_style)

    for item in items:
        if item.section == "annex-I-technical-specification":
            add_blank_body_placeholder(doc, normal_style)
            continue

        child_title = titleize_slug(item.short_name)
        add_styled_paragraph(doc, child_title, h2_style)
        add_blank_body_placeholder(doc, normal_style)


def add_back_matter(
    doc: Document,
    h1_style: str,
    normal_style: str,
) -> None:
    """
    Add standard closing sections, each starting on a new page.
    """
    add_page_break(doc)
    add_styled_paragraph(doc, "REFERENCES", h1_style)
    add_blank_body_placeholder(doc, normal_style)

    add_page_break(doc)
    add_styled_paragraph(doc, "LIST OF ABBREVIATIONS", h1_style)
    add_blank_body_placeholder(doc, normal_style)

    add_page_break(doc)
    add_styled_paragraph(doc, "CONTRIBUTORS TO DRAFTING AND REVIEW", h1_style)
    add_blank_body_placeholder(doc, normal_style)


# ---------------------------------------------------------------------------
# Main workflow
# ---------------------------------------------------------------------------

def main() -> int:
    """
    Create a skeleton TECDOC master by cloning the template and rebuilding only
    the structural pages and headings.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Template not found: {TEMPLATE_PATH}")

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)

    # Copy the template so the output inherits the original styles and setup.
    shutil.copy2(TEMPLATE_PATH, OUTPUT_PATH)

    # Open the copied file and rebuild its body.
    doc = Document(OUTPUT_PATH)

    # Find styles that exist in the template.
    h1_style = find_first_existing_style(doc, ["H1", "Heading 1", "HEADING1"])
    h2_style = find_first_existing_style(doc, ["H2", "Heading 2", "HEADING2"])
    normal_style = find_first_existing_style(doc, ["Normal", "normal"])

    # Front matter styles. Fall back gently if the template does not define
    # dedicated title/subtitle styles.
    tecdoc_style = find_first_existing_style(
        doc,
        ["Subtitle", "subtitle", "Normal", "normal"],
        fallback="Normal",
    )
    title_style = find_first_existing_style(
        doc,
        ["Title", "title", "H1", "Heading 1"],
        fallback=h1_style,
    )
    subtitle_style = find_first_existing_style(
        doc,
        ["Subtitle", "subtitle", "Normal", "normal"],
        fallback=normal_style,
    )

    # Scan available source files.
    chapters = scan_source_items(SOURCE_ROOT / "chapters", "chapters")
    annex1 = scan_source_items(
        SOURCE_ROOT / "annex-I-technical-specification",
        "annex-I-technical-specification",
    )
    annex2 = scan_source_items(
        SOURCE_ROOT / "annex-II-organizations",
        "annex-II-organizations",
    )
    annex3 = scan_source_items(
        SOURCE_ROOT / "annex-III-codes",
        "annex-III-codes",
    )
    annex4 = scan_source_items(
        SOURCE_ROOT / "annex-IV-individual-results",
        "annex-IV-individual-results",
    )

    # Clear sample content but preserve styles and section/page settings.
    clear_document_body_preserve_sections(doc)

    # Front matter
    add_title_page(
        doc,
        tecdoc_style=tecdoc_style,
        title_style=title_style,
        subtitle_style=subtitle_style,
    )

    add_page_break(doc)
    add_foreword_page(doc, h1_style=h1_style, normal_style=normal_style)

    add_page_break(doc)
    add_toc_page(doc, h1_style=h1_style, normal_style=normal_style)

    # Main body
    if chapters:
        add_page_break(doc)
        add_chapter_skeleton(
            doc,
            chapters=chapters,
            h1_style=h1_style,
            normal_style=normal_style,
        )

    add_annex_group(
        doc,
        group_title="ANNEX I. TECHNICAL SPECIFICATION",
        items=annex1,
        h1_style=h1_style,
        h2_style=h2_style,
        normal_style=normal_style,
    )
    add_annex_group(
        doc,
        group_title="ANNEX II. DESCRIPTION OF ORGANIZATIONS",
        items=annex2,
        h1_style=h1_style,
        h2_style=h2_style,
        normal_style=normal_style,
    )
    add_annex_group(
        doc,
        group_title="ANNEX III. DESCRIPTION OF CODES",
        items=annex3,
        h1_style=h1_style,
        h2_style=h2_style,
        normal_style=normal_style,
    )
    add_annex_group(
        doc,
        group_title="ANNEX IV. INDIVIDUAL RESULTS",
        items=annex4,
        h1_style=h1_style,
        h2_style=h2_style,
        normal_style=normal_style,
    )

    add_back_matter(doc, h1_style=h1_style, normal_style=normal_style)

    doc.save(OUTPUT_PATH)

    print(f"Created skeleton: {OUTPUT_PATH}")
    print(
        "Styles used: "
        f"Title={title_style}, Subtitle={subtitle_style}, "
        f"H1={h1_style}, H2={h2_style}, Normal={normal_style}"
    )
    return 0


if __name__ == "__main__":
    raise SystemExit(main())