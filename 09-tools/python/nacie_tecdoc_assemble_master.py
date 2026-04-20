#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path
from typing import Optional

import yaml
from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.table import Table


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

SCRIPT_PATH = Path(__file__).resolve()
ROOT = SCRIPT_PATH.parents[2]

DEFAULT_SKELETON_PATH = ROOT / "06-master" / "01-nacie_tecdoc_master_skeleton.docx"
DEFAULT_OUTPUT_PATH = ROOT / "05-build" / "02-nacie_tecdoc_master_chapter03_inserted.docx"


# ---------------------------------------------------------------------------
# Low-level DOCX helpers
# ---------------------------------------------------------------------------

def iter_block_items(parent: _Document):
    """
    Yield paragraphs and tables from a document body in order.
    """
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    style_name: Optional[str] = None,
) -> Paragraph:
    """
    Insert a new paragraph directly after an existing paragraph.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    if text:
        new_para.add_run(text)

    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass

    return new_para


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def find_first_existing_style(
    doc: Document,
    candidates: list[str],
    fallback: str = "Normal",
) -> str:
    """
    Return the first style name that exists in the document.
    """
    available = {style.name for style in doc.styles}
    for name in candidates:
        if name in available:
            return name
    return fallback


# ---------------------------------------------------------------------------
# YAML helpers
# ---------------------------------------------------------------------------

def load_yaml(path: Path) -> dict:
    """
    Load YAML from disk.
    """
    with path.open("r", encoding="utf-8") as handle:
        return yaml.safe_load(handle)


# ---------------------------------------------------------------------------
# Heading matching helpers
# ---------------------------------------------------------------------------

def normalize_heading_text(text: str) -> str:
    """
    Normalize heading text for matching.
    """
    return " ".join(text.strip().lower().split())


def is_references_heading(text: str) -> bool:
    """
    Return True if a heading text marks a references section.
    """
    normalized = normalize_heading_text(text)
    return normalized in {"references", "reference", "bibliography"}


def find_heading_paragraph(doc: Document, heading_text: str) -> Paragraph:
    """
    Find the first paragraph whose text matches the requested heading.
    """
    target = normalize_heading_text(heading_text)

    for paragraph in doc.paragraphs:
        if normalize_heading_text(paragraph.text) == target:
            return paragraph

    raise ValueError(f"Heading not found in skeleton: {heading_text!r}")


# ---------------------------------------------------------------------------
# Cleanup helpers
# ---------------------------------------------------------------------------

def remove_placeholder_paragraphs_after_heading(heading_para: Paragraph) -> Paragraph:
    """
    Remove the default placeholder paragraph(s) that follow a section heading.
    """
    parent = heading_para._parent

    while True:
        next_el = heading_para._p.getnext()
        if next_el is None or not next_el.tag.endswith("}p"):
            break

        next_para = Paragraph(next_el, parent)
        next_text = (next_para.text or "").strip()

        if next_text == "Body text placeholder.":
            next_el.getparent().remove(next_el)
            continue

        if next_text == "":
            next_el.getparent().remove(next_el)
            continue

        break

    return heading_para


# ---------------------------------------------------------------------------
# Block insertion helpers
# ---------------------------------------------------------------------------

def insert_paragraph_block(
    anchor: Paragraph,
    text: str,
    style_name: str,
) -> Paragraph:
    """
    Insert one paragraph block after the anchor.

    Inline fields like {REF:...}, {FIG:...}, {TAB:...} are preserved exactly.
    """
    return insert_paragraph_after(anchor, text=text, style_name=style_name)


def insert_heading_block(
    anchor: Paragraph,
    text: str,
    level: int,
    h1_style: str,
    h2_style: str,
    h3_style: str,
) -> Paragraph:
    """
    Insert one heading block after the anchor.
    """
    if level == 1:
        style_name = h1_style
    elif level == 2:
        style_name = h2_style
    else:
        style_name = h3_style

    return insert_paragraph_after(anchor, text=text, style_name=style_name)


# ---------------------------------------------------------------------------
# Reference insertion
# ---------------------------------------------------------------------------

def append_references_to_master(
    doc: Document,
    references: list[dict],
    normal_style: str,
) -> None:
    """
    Append real reference texts to the global REFERENCES chapter in the form:

      {REF:chapter03_local_001} Reference text...

    Keep a page break after the inserted references block so the next major
    section still starts on a new page.
    """
    if not references:
        return

    references_heading = find_heading_paragraph(doc, "REFERENCES")
    anchor = remove_placeholder_paragraphs_after_heading(references_heading)

    for ref in references:
        ref_id = ref.get("id", "REF:unknown")
        raw_text = (ref.get("raw_text") or "").strip()
        if not raw_text:
            continue

        line = "{" + ref_id + "} " + raw_text
        anchor = insert_paragraph_after(anchor, text=line, style_name=normal_style)

    # Preserve separation from the next major section.
    page_break_para = insert_paragraph_after(anchor, style_name=normal_style)
    run = page_break_para.add_run()
    run.add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Assembly logic
# ---------------------------------------------------------------------------

def assemble_one_document(
    skeleton_path: Path,
    content_yaml_path: Path,
    output_path: Path,
) -> None:
    """
    Insert one normalized document into the skeleton and save the result.

    Temporary verification behaviour:
    - insert headings and paragraphs only
    - preserve all inline fields exactly
    - do NOT insert standalone figure/table markers
    - do NOT insert a local References subsection inside the chapter
    - append bibliography entries to the global REFERENCES chapter
    """
    if not skeleton_path.exists():
        raise FileNotFoundError(f"Skeleton file not found: {skeleton_path}")

    if not content_yaml_path.exists():
        raise FileNotFoundError(f"Normalized content YAML not found: {content_yaml_path}")

    payload = load_yaml(content_yaml_path)

    title = payload.get("title")
    if not title:
        raise ValueError("Normalized content is missing 'title'.")

    blocks = payload.get("blocks", [])
    if not isinstance(blocks, list):
        raise ValueError("'blocks' must be a list in content.yaml")

    references = payload.get("references", [])
    if not isinstance(references, list):
        raise ValueError("'references' must be a list in content.yaml")

    doc = Document(skeleton_path)

    h1_style = find_first_existing_style(doc, ["H1", "Heading 1", "HEADING1"])
    h2_style = find_first_existing_style(doc, ["H2", "Heading 2", "HEADING2"], fallback=h1_style)
    h3_style = find_first_existing_style(doc, ["H3", "Heading 3", "HEADING3"], fallback=h2_style)
    normal_style = find_first_existing_style(doc, ["Normal", "normal"])

    # Insert main body content under the chapter heading.
    heading_para = find_heading_paragraph(doc, title)
    anchor = remove_placeholder_paragraphs_after_heading(heading_para)

    for block in blocks:
        block_type = block.get("type")

        if block_type == "heading":
            text = block.get("text", "")
            # Do not insert a local "References" heading inside the chapter.
            if is_references_heading(text):
                continue

            anchor = insert_heading_block(
                anchor=anchor,
                text=text,
                level=int(block.get("level", 1)),
                h1_style=h1_style,
                h2_style=h2_style,
                h3_style=h3_style,
            )
            continue

        if block_type == "paragraph":
            anchor = insert_paragraph_block(
                anchor=anchor,
                text=block.get("text", ""),
                style_name=normal_style,
            )
            continue

        if block_type == "figure":
            # Ignore standalone figure blocks for now.
            continue

        if block_type == "table":
            # Ignore standalone table blocks for now.
            continue

        anchor = insert_paragraph_block(
            anchor=anchor,
            text=f"[Unsupported block type: {block_type}]",
            style_name=normal_style,
        )

    # Preserve separation from the next chapter after inserted body content.
    page_break_para = insert_paragraph_after(anchor, style_name=normal_style)
    run = page_break_para.add_run()
    run.add_break(WD_BREAK.PAGE)

    # Append the real bibliography texts to the global REFERENCES chapter.
    append_references_to_master(
        doc=doc,
        references=references,
        normal_style=normal_style,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    """
    Parse command-line arguments.
    """
    parser = argparse.ArgumentParser(
        description="Insert one normalized NACIE TECDOC document into the master skeleton."
    )
    parser.add_argument(
        "content_yaml",
        help="Path to content.yaml produced by nacie_tecdoc_normalize_docx.py",
    )
    parser.add_argument(
        "--skeleton",
        default=str(DEFAULT_SKELETON_PATH),
        help="Path to the skeleton master .docx",
    )
    parser.add_argument(
        "--output",
        default=str(DEFAULT_OUTPUT_PATH),
        help="Path to output assembled .docx",
    )
    return parser.parse_args()


def main() -> int:
    """
    Run the insertion prototype.
    """
    args = parse_args()

    content_yaml_path = Path(args.content_yaml).resolve()
    skeleton_path = Path(args.skeleton).resolve()
    output_path = Path(args.output).resolve()

    assemble_one_document(
        skeleton_path=skeleton_path,
        content_yaml_path=content_yaml_path,
        output_path=output_path,
    )

    print(f"Skeleton input : {skeleton_path}")
    print(f"Content input  : {content_yaml_path}")
    print(f"Output written : {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())