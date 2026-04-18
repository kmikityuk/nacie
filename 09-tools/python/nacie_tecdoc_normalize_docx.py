#!/usr/bin/env python3
from __future__ import annotations

import argparse
import hashlib
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

import yaml
from docx import Document


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

SCRIPT_PATH = Path(__file__).resolve()
ROOT = SCRIPT_PATH.parents[2]

DEFAULT_SOURCE_ROOT = ROOT / "01-source"
DEFAULT_NORMALIZED_ROOT = ROOT / "03-normalized"
DEFAULT_ASSETS_ROOT = ROOT / "04-assets"


# ---------------------------------------------------------------------------
# Data model
# ---------------------------------------------------------------------------

@dataclass
class SourceInfo:
    """
    Parsed information about one source document path.
    """
    source_path: Path
    section: str
    item_no: Optional[int]
    short_name: str
    version: Optional[str]
    document_id: str
    title: str


# ---------------------------------------------------------------------------
# Path and naming helpers
# ---------------------------------------------------------------------------

def classify_section(source_path: Path, source_root: Path) -> str:
    """
    Determine the logical section based on the path relative to 01-source.
    """
    rel_path = source_path.relative_to(source_root)
    if len(rel_path.parts) < 2:
        return "unknown"

    top = rel_path.parts[0]
    if top == "chapters":
        return "chapters"
    if top == "annex-I-technical-specification":
        return "annex-I-technical-specification"
    if top == "annex-II-organizations":
        return "annex-II-organizations"
    if top == "annex-III-codes":
        return "annex-III-codes"
    if top == "annex-IV-individual-results":
        return "annex-IV-individual-results"

    return top


def section_to_output_subdir(section: str) -> str:
    """
    Map source section names to normalized/assets subdirectories.
    """
    mapping = {
        "chapters": "chapters",
        "annex-I-technical-specification": "annex-I-technical-specification",
        "annex-II-organizations": "annex-II-organizations",
        "annex-III-codes": "annex-III-codes",
        "annex-IV-individual-results": "annex-IV-individual-results",
    }
    return mapping.get(section, section)


def parse_source_filename(path: Path) -> tuple[Optional[int], str, Optional[str]]:
    """
    Parse filenames such as:
      03-chapter03-v01.docx
      11-KIT-v01.docx
      annexI-technical-specification-v00.docx
    """
    name = path.name

    pattern_with_item = re.compile(
        r"^(?P<item>\d{2})-(?P<short>.+)-v(?P<ver>\d{2})\.(docx|pdf|txt|md)$",
        re.IGNORECASE,
    )
    pattern_without_item = re.compile(
        r"^(?P<short>.+)-v(?P<ver>\d{2})\.(docx|pdf|txt|md)$",
        re.IGNORECASE,
    )

    match = pattern_with_item.fullmatch(name)
    if match:
        return int(match.group("item")), match.group("short"), match.group("ver")

    match = pattern_without_item.fullmatch(name)
    if match:
        return None, match.group("short"), match.group("ver")

    return None, path.stem, None


def slugify(value: str) -> str:
    """
    Convert text into a conservative filesystem-safe slug.
    """
    value = value.strip().replace(" ", "-")
    value = re.sub(r"[^A-Za-z0-9._-]+", "-", value)
    value = re.sub(r"-{2,}", "-", value)
    return value.strip("-").lower()


def build_document_id(section: str, item_no: Optional[int], short_name: str) -> str:
    """
    Build a stable internal document identifier.

    Examples:
      chapter03
      annexii_11_kit
      annexiii_21_sam-moose-scm
      annexiv_01_anl
      annexi_technical_specification
    """
    short_slug = slugify(short_name)

    if section == "chapters":
        if item_no is not None:
            return f"chapter{item_no:02d}"
        return f"chapter_{short_slug}"

    if section == "annex-I-technical-specification":
        return "annexi_technical_specification"

    if section == "annex-II-organizations":
        if item_no is not None:
            return f"annexii_{item_no:02d}_{short_slug}"
        return f"annexii_{short_slug}"

    if section == "annex-III-codes":
        if item_no is not None:
            return f"annexiii_{item_no:02d}_{short_slug}"
        return f"annexiii_{short_slug}"

    if section == "annex-IV-individual-results":
        if item_no is not None:
            return f"annexiv_{item_no:02d}_{short_slug}"
        return f"annexiv_{short_slug}"

    return short_slug


def titleize_slug(slug: str) -> str:
    """
    Convert filename slug into a readable title.
    """
    special = {
        "chapter01": "Chapter 1",
        "chapter02": "Chapter 2",
        "chapter03": "Chapter 3",
        "chapter04": "Chapter 4",
        "chapter05": "Chapter 5",
        "chapter06": "Chapter 6",
        "chapter07": "Chapter 7",
        "chapter08": "Chapter 8",
        "chapter09": "Chapter 9",
        "conclusions": "Conclusions",
        "references": "References",
        "abbreviations": "List of Abbreviations",
        "annexI-technical-specification": "Annex I. Technical Specification",
    }
    if slug in special:
        return special[slug]

    return slug.replace("-", " ").replace("_", " ")


def build_source_info(source_path: Path, source_root: Path) -> SourceInfo:
    """
    Parse one source path into structured source information.
    """
    section = classify_section(source_path, source_root)
    item_no, short_name, version = parse_source_filename(source_path)
    document_id = build_document_id(section, item_no, short_name)
    title = titleize_slug(short_name)

    return SourceInfo(
        source_path=source_path,
        section=section,
        item_no=item_no,
        short_name=short_name,
        version=version,
        document_id=document_id,
        title=title,
    )


# ---------------------------------------------------------------------------
# General helpers
# ---------------------------------------------------------------------------

def sha256_file(path: Path) -> str:
    """
    Compute SHA256 for the source file, useful for traceability.
    """
    h = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def ensure_dir(path: Path) -> None:
    """
    Create a directory if needed.
    """
    path.mkdir(parents=True, exist_ok=True)


def load_docx(source_path: Path) -> Document:
    """
    Open the DOCX file with python-docx.
    """
    return Document(source_path)


def get_paragraph_style_name(paragraph) -> str:
    """
    Return the paragraph style name if present.
    """
    try:
        if paragraph.style is not None and paragraph.style.name:
            return paragraph.style.name
    except Exception:
        pass
    return ""


def infer_heading_level(style_name: str, paragraph_text: str) -> Optional[int]:
    """
    Infer a heading level from the paragraph style and text.

    This is intentionally conservative:
    - H1 / Heading 1 -> level 1
    - H2 / Heading 2 -> level 2
    - H3 / Heading 3 -> level 3
    """
    style = style_name.strip().lower()

    if style in {"h1", "heading 1", "heading1"}:
        return 1
    if style in {"h2", "heading 2", "heading2"}:
        return 2
    if style in {"h3", "heading 3", "heading3"}:
        return 3

    _ = paragraph_text
    return None


def is_placeholder_source(info: SourceInfo, source_path: Path) -> bool:
    """
    Detect whether the source appears to be a placeholder.
    """
    size_bytes = source_path.stat().st_size
    if info.version == "00":
        return True
    if size_bytes == 13325:
        return True
    return False


def is_references_heading(text: str) -> bool:
    """
    Return True if a heading text marks the start of a references section.
    """
    normalized = text.strip().lower()
    return normalized in {
        "references",
        "reference",
        "bibliography",
    }


# ---------------------------------------------------------------------------
# Inline marker extraction
# ---------------------------------------------------------------------------

_REF_NUM_PATTERN = re.compile(r"\[(\d+)\]")


def paragraph_text_with_markers(
    text: str,
    local_ref_map: dict[str, str],
    document_id: str,
) -> str:
    """
    Convert some obvious local citation markers into internal placeholders.

    This is intentionally conservative. It only rewrites references like [7]
    into {REF:<document_id>_local_007}. Figures/tables are not rewritten
    automatically here because in many source documents they are ambiguous and
    require spatial context.
    """
    def repl_ref(match: re.Match[str]) -> str:
        num = match.group(1)
        key = f"REF:{document_id}_local_{int(num):03d}"
        local_ref_map[num] = key
        return "{" + key + "}"

    return _REF_NUM_PATTERN.sub(repl_ref, text)


# ---------------------------------------------------------------------------
# Image extraction
# ---------------------------------------------------------------------------

def extract_images_from_docx(
    source_path: Path,
    images_dir: Path,
    document_id: str,
) -> list[dict]:
    """
    Extract embedded images from the DOCX zip package.

    Returns
    -------
    list[dict]
        Metadata for extracted image assets in source order as far as the ZIP
        layout allows. In v1 we do not map exact image positions to paragraphs.
    """
    ensure_dir(images_dir)

    image_records: list[dict] = []

    with zipfile.ZipFile(source_path, "r") as zf:
        media_names = sorted(
            [name for name in zf.namelist() if name.startswith("word/media/")]
        )

        for index, media_name in enumerate(media_names, start=1):
            suffix = Path(media_name).suffix.lower() or ".bin"
            image_name = f"{document_id}_fig_{index:03d}{suffix}"
            out_path = images_dir / image_name

            with zf.open(media_name) as src, out_path.open("wb") as dst:
                dst.write(src.read())

            image_records.append(
                {
                    "asset_name": image_name,
                    "source_zip_path": media_name,
                    "relative_path": f"images/{image_name}",
                }
            )

    return image_records


# ---------------------------------------------------------------------------
# Table extraction
# ---------------------------------------------------------------------------

def extract_tables(
    doc: Document,
    tables_dir: Path,
    document_id: str,
) -> list[dict]:
    """
    Extract Word tables into separate YAML files.

    Returns
    -------
    list[dict]
        Metadata for normalized table assets.
    """
    ensure_dir(tables_dir)

    table_records: list[dict] = []

    for index, table in enumerate(doc.tables, start=1):
        rows = []
        for row in table.rows:
            row_values = []
            for cell in row.cells:
                row_values.append(cell.text.strip())
            rows.append(row_values)

        table_name = f"{document_id}_table_{index:03d}.yaml"
        out_path = tables_dir / table_name

        with out_path.open("w", encoding="utf-8") as handle:
            yaml.safe_dump(
                {
                    "rows": rows,
                },
                handle,
                sort_keys=False,
                allow_unicode=True,
            )

        table_records.append(
            {
                "table_index": index,
                "file": f"tables/{table_name}",
                "rows": len(rows),
                "cols": max((len(r) for r in rows), default=0),
            }
        )

    return table_records


# ---------------------------------------------------------------------------
# Reference extraction
# ---------------------------------------------------------------------------

def extract_reference_entries(doc: Document, document_id: str) -> list[dict]:
    """
    Extract bibliography entries from either:
    1. paragraphs that begin with [n]
    2. paragraphs located under a heading such as 'References'

    This is a conservative first pass:
    - numbered references keep their numeric order if present
    - plain reference paragraphs inside a References section are also captured
    - entries are deduplicated by raw_text
    """
    refs: list[dict] = []

    in_references_section = False
    local_counter = 0

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = get_paragraph_style_name(paragraph)

        if not text:
            continue

        heading_level = infer_heading_level(style_name, text)

        if heading_level is not None and is_references_heading(text):
            in_references_section = True
            continue

        if heading_level is not None and in_references_section:
            in_references_section = False

        match = re.match(r"^\[(\d+)\]\s*(.+)$", text)
        if match:
            num = int(match.group(1))
            body = match.group(2).strip()
            ref_id = f"REF:{document_id}_local_{num:03d}"
            refs.append(
                {
                    "id": ref_id,
                    "raw_text": body,
                    "source_label": f"[{num}]",
                }
            )
            continue

        if in_references_section:
            local_counter += 1
            ref_id = f"REF:{document_id}_local_{local_counter:03d}"
            refs.append(
                {
                    "id": ref_id,
                    "raw_text": text,
                    "source_label": None,
                }
            )

    deduped: list[dict] = []
    seen_texts: set[str] = set()

    for ref in refs:
        key = ref["raw_text"].strip()
        if key in seen_texts:
            continue
        deduped.append(ref)
        seen_texts.add(key)

    return deduped


# ---------------------------------------------------------------------------
# Block extraction
# ---------------------------------------------------------------------------

def extract_blocks(
    doc: Document,
    info: SourceInfo,
    image_records: list[dict],
    table_records: list[dict],
    warnings: list[str],
) -> list[dict]:
    """
    Extract a conservative ordered block list from the document.

    v1 behaviour:
    - headings from paragraph styles
    - paragraphs from non-empty text paragraphs
    - paragraphs inside a References section are excluded from blocks and stored
      only in the references list
    - append figure blocks after textual extraction if images exist
    - append table blocks after textual extraction if tables exist

    This does not yet reconstruct exact interleaving of figures/tables with
    surrounding paragraphs. That refinement can come in v2.
    """
    blocks: list[dict] = []
    local_ref_map: dict[str, str] = {}
    in_references_section = False

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        style_name = get_paragraph_style_name(paragraph)

        if not text:
            continue

        heading_level = infer_heading_level(style_name, text)

        if heading_level is not None and is_references_heading(text):
            in_references_section = True
            blocks.append(
                {
                    "type": "heading",
                    "level": heading_level,
                    "text": text,
                    "style": style_name or None,
                }
            )
            continue

        if heading_level is not None and in_references_section:
            in_references_section = False

        if heading_level is not None:
            blocks.append(
                {
                    "type": "heading",
                    "level": heading_level,
                    "text": text,
                    "style": style_name or None,
                }
            )
            continue

        match = re.match(r"^\[(\d+)\]\s*(.+)$", text)
        if match:
            # Numbered bibliography entry: keep it out of body blocks.
            continue

        if in_references_section:
            # Plain bibliography entry under a References heading: keep it out of
            # body blocks, it will go to references: only.
            continue

        normalized_text = paragraph_text_with_markers(
            text=text,
            local_ref_map=local_ref_map,
            document_id=info.document_id,
        )
        blocks.append(
            {
                "type": "paragraph",
                "text": normalized_text,
                "style": style_name or None,
            }
        )

    for index, table_meta in enumerate(table_records, start=1):
        table_id = f"TAB:{info.document_id}_{index:03d}"
        blocks.append(
            {
                "type": "table",
                "id": table_id,
                "caption": "",
                "file": table_meta["file"],
            }
        )
        warnings.append(
            f"Table {table_id} extracted without automatic caption detection; caption must be reviewed manually."
        )

    for index, image_meta in enumerate(image_records, start=1):
        fig_id = f"FIG:{info.document_id}_{index:03d}"
        blocks.append(
            {
                "type": "figure",
                "id": fig_id,
                "image": image_meta["relative_path"],
                "caption": "",
            }
        )
        warnings.append(
            f"Figure {fig_id} extracted without automatic caption detection; caption must be reviewed manually."
        )

    if image_records:
        warnings.append(
            "In-text figure references are not automatically resolved in v1; review manually."
        )
    if table_records:
        warnings.append(
            "In-text table references are not automatically resolved in v1; review manually."
        )

    return blocks


# ---------------------------------------------------------------------------
# YAML writing
# ---------------------------------------------------------------------------

def write_content_yaml(
    out_path: Path,
    info: SourceInfo,
    source_root: Path,
    source_hash: str,
    blocks: list[dict],
    references: list[dict],
    placeholder: bool,
) -> None:
    """
    Write the main normalized content YAML.
    """
    payload = {
        "document_id": info.document_id,
        "source_path": str(info.source_path.relative_to(source_root)),
        "section": info.section,
        "item_no": info.item_no,
        "version": info.version,
        "title": info.title,
        "status": "placeholder" if placeholder else "present",
        "source_hash": source_hash,
        "blocks": blocks,
        "references": references,
    }

    with out_path.open("w", encoding="utf-8") as handle:
        yaml.safe_dump(
            payload,
            handle,
            sort_keys=False,
            allow_unicode=True,
            width=100,
        )


def write_warnings_yaml(out_path: Path, warnings: list[str]) -> None:
    """
    Write normalization warnings.
    """
    payload = {"warnings": warnings}
    with out_path.open("w", encoding="utf-8") as handle:
        yaml.safe_dump(
            payload,
            handle,
            sort_keys=False,
            allow_unicode=True,
            width=100,
        )


# ---------------------------------------------------------------------------
# Main workflow
# ---------------------------------------------------------------------------

def normalize_one_docx(
    source_path: Path,
    source_root: Path,
    normalized_root: Path,
    assets_root: Path,
) -> tuple[Path, Path]:
    """
    Normalize one DOCX source into YAML + assets.

    Returns
    -------
    tuple[Path, Path]
        (content_yaml_path, warnings_yaml_path)
    """
    if not source_path.exists():
        raise FileNotFoundError(f"Source file not found: {source_path}")

    info = build_source_info(source_path, source_root)
    source_hash = sha256_file(source_path)

    output_subdir = section_to_output_subdir(info.section)

    normalized_dir = normalized_root / output_subdir / info.document_id
    assets_dir = assets_root / output_subdir / info.document_id
    images_dir = assets_dir / "images"
    tables_dir = assets_dir / "tables"

    ensure_dir(normalized_dir)
    ensure_dir(images_dir)
    ensure_dir(tables_dir)

    warnings: list[str] = []
    placeholder = is_placeholder_source(info, source_path)

    if placeholder:
        warnings.append("Source file appears to be a placeholder.")

    doc = load_docx(source_path)

    image_records = extract_images_from_docx(
        source_path=source_path,
        images_dir=images_dir,
        document_id=info.document_id,
    )
    table_records = extract_tables(
        doc=doc,
        tables_dir=tables_dir,
        document_id=info.document_id,
    )

    blocks = extract_blocks(
        doc=doc,
        info=info,
        image_records=image_records,
        table_records=table_records,
        warnings=warnings,
    )

    references = extract_reference_entries(
        doc=doc,
        document_id=info.document_id,
    )

    if not blocks and not placeholder:
        warnings.append("No content blocks were extracted from a non-placeholder source.")
    if not references:
        warnings.append("No explicit bibliography entries were extracted.")
    if image_records and not any(block["type"] == "figure" for block in blocks):
        warnings.append("Images were extracted but no figure blocks were created.")
    if table_records and not any(block["type"] == "table" for block in blocks):
        warnings.append("Tables were extracted but no table blocks were created.")

    content_yaml_path = normalized_dir / "content.yaml"
    warnings_yaml_path = normalized_dir / "warnings.yaml"

    write_content_yaml(
        out_path=content_yaml_path,
        info=info,
        source_root=source_root,
        source_hash=source_hash,
        blocks=blocks,
        references=references,
        placeholder=placeholder,
    )
    write_warnings_yaml(warnings_yaml_path, warnings)

    return content_yaml_path, warnings_yaml_path


def parse_args() -> argparse.Namespace:
    """
    Parse command-line arguments.
    """
    parser = argparse.ArgumentParser(
        description="Normalize one NACIE TECDOC DOCX source into YAML plus extracted assets."
    )
    parser.add_argument(
        "source_docx",
        help="Path to the source .docx file to normalize.",
    )
    parser.add_argument(
        "--source-root",
        default=str(DEFAULT_SOURCE_ROOT),
        help="Root directory for source documents. Default: project 01-source.",
    )
    parser.add_argument(
        "--normalized-root",
        default=str(DEFAULT_NORMALIZED_ROOT),
        help="Output root for normalized YAML packages. Default: project 03-normalized.",
    )
    parser.add_argument(
        "--assets-root",
        default=str(DEFAULT_ASSETS_ROOT),
        help="Output root for extracted assets. Default: project 04-assets.",
    )
    return parser.parse_args()


def main() -> int:
    """
    Run the normalizer for one source document.
    """
    args = parse_args()

    source_path = Path(args.source_docx).resolve()
    source_root = Path(args.source_root).resolve()
    normalized_root = Path(args.normalized_root).resolve()
    assets_root = Path(args.assets_root).resolve()

    content_yaml_path, warnings_yaml_path = normalize_one_docx(
        source_path=source_path,
        source_root=source_root,
        normalized_root=normalized_root,
        assets_root=assets_root,
    )

    print(f"Normalized source: {source_path}")
    print(f"content.yaml : {content_yaml_path}")
    print(f"warnings.yaml: {warnings_yaml_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())