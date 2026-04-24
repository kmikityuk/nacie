#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path
from typing import Optional

import yaml
from docx import Document
from docx.document import Document as _Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx.text.paragraph import Paragraph
from docx.table import Table


# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

SCRIPT_PATH = Path(__file__).resolve()
ROOT = SCRIPT_PATH.parents[2]

DEFAULT_SKELETON_PATH = ROOT / "06-master" / "01-nacie_tecdoc_master_skeleton.docx"
DEFAULT_OUTPUT_PATH = ROOT / "05-build" / "02-nacie_tecdoc_master_chapter03_inserted.docx"


# ---------------------------------------------------------------------------
# Low-level DOCX helpers
# ---------------------------------------------------------------------------

def iter_block_items(parent: _Document):
    """
    Yield paragraphs and tables from a document body in order.
    """
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def insert_paragraph_after(
    paragraph: Paragraph,
    text: str = "",
    style_name: Optional[str] = None,
) -> Paragraph:
    """
    Insert a new paragraph directly after an existing paragraph.
    """
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    if text:
        new_para.add_run(text)

    if style_name:
        try:
            new_para.style = style_name
        except Exception:
            pass

    return new_para


def insert_table_after(paragraph: Paragraph, rows: list[list[str]]) -> tuple[Table, Paragraph]:
    """
    Insert a Word table directly after an existing paragraph.

    Returns
    -------
    tuple[Table, Paragraph]
        The inserted table and a trailing paragraph inserted after the table,
        which can be used as the next insertion anchor.
    """
    if not rows:
        rows = [[""]]

    n_rows = len(rows)
    n_cols = max((len(r) for r in rows), default=1)

    table = paragraph._parent.add_table(rows=n_rows, cols=n_cols, width=Inches(6.0))
    tbl_element = table._tbl

    for r_idx, row in enumerate(rows):
        for c_idx in range(n_cols):
            value = row[c_idx] if c_idx < len(row) else ""
            table.cell(r_idx, c_idx).text = value

    paragraph._p.addnext(tbl_element)

    trailing_p = OxmlElement("w:p")
    tbl_element.addnext(trailing_p)
    trailing_para = Paragraph(trailing_p, paragraph._parent)

    return table, trailing_para


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    """
    Replace all runs in a paragraph with plain text, preserving the paragraph style.
    """
    p = paragraph._p
    for child in list(p):
        if child.tag.endswith("}r"):
            p.remove(child)

    paragraph.add_run(text)


def insert_page_break_after(paragraph: Paragraph, style_name: Optional[str] = None) -> Paragraph:
    """
    Insert a page-break paragraph directly after an existing paragraph.
    """
    page_para = insert_paragraph_after(paragraph, style_name=style_name)
    run = page_para.add_run()
    run.add_break(WD_BREAK.PAGE)
    return page_para


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------

def find_first_existing_style(
    doc: Document,
    candidates: list[str],
    fallback: str = "Normal",
) -> str:
    """
    Return the first style name that exists in the document.
    """
    available = {style.name for style in doc.styles}
    for name in candidates:
        if name in available:
            return name
    return fallback


# ---------------------------------------------------------------------------
# YAML helpers
# ---------------------------------------------------------------------------

def load_yaml(path: Path) -> dict:
    """
    Load YAML from disk.
    """
    with path.open("r", encoding="utf-8") as handle:
        return yaml.safe_load(handle)


def load_table_rows(table_yaml_path: Path) -> list[list[str]]:
    """
    Load extracted table rows from a YAML file.
    """
    payload = load_yaml(table_yaml_path)
    rows = payload.get("rows", [])
    if not isinstance(rows, list):
        return []
    return rows


# ---------------------------------------------------------------------------
# Heading matching helpers
# ---------------------------------------------------------------------------

def normalize_heading_text(text: str) -> str:
    """
    Normalize heading text for matching.
    """
    return " ".join(text.strip().lower().split())


def is_references_heading(text: str) -> bool:
    """
    Return True if a heading text marks a references section.
    """
    normalized = normalize_heading_text(text)
    return normalized in {"references", "reference", "bibliography"}


def find_heading_paragraph(doc: Document, heading_text: str) -> Paragraph:
    """
    Find the first paragraph whose text matches the requested heading.
    """
    target = normalize_heading_text(heading_text)

    for paragraph in doc.paragraphs:
        if normalize_heading_text(paragraph.text) == target:
            return paragraph

    raise ValueError(f"Heading not found in skeleton: {heading_text!r}")


# ---------------------------------------------------------------------------
# Cleanup helpers
# ---------------------------------------------------------------------------

def remove_placeholder_paragraphs_after_heading(heading_para: Paragraph) -> Paragraph:
    """
    Remove the default placeholder paragraph(s) that follow a section heading.
    """
    parent = heading_para._parent

    while True:
        next_el = heading_para._p.getnext()
        if next_el is None or not next_el.tag.endswith("}p"):
            break

        next_para = Paragraph(next_el, parent)
        next_text = (next_para.text or "").strip()

        if next_text == "Body text placeholder.":
            next_el.getparent().remove(next_el)
            continue

        if next_text == "":
            next_el.getparent().remove(next_el)
            continue

        break

    return heading_para


# ---------------------------------------------------------------------------
# Block insertion helpers
# ---------------------------------------------------------------------------

def insert_paragraph_block(
    anchor: Paragraph,
    text: str,
    style_name: str,
) -> Paragraph:
    """
    Insert one paragraph block after the anchor.

    Inline fields like {REF:...}, {FIG_REF:...}, {TAB_REF:...} are preserved exactly.
    """
    return insert_paragraph_after(anchor, text=text, style_name=style_name)


def insert_heading_block(
    anchor: Paragraph,
    text: str,
    level: int,
    h1_style: str,
    h2_style: str,
    h3_style: str,
) -> Paragraph:
    """
    Insert one heading block after the anchor.
    """
    if level == 1:
        style_name = h1_style
    elif level == 2:
        style_name = h2_style
    else:
        style_name = h3_style

    return insert_paragraph_after(anchor, text=text, style_name=style_name)


# ---------------------------------------------------------------------------
# Reference insertion
# ---------------------------------------------------------------------------

def append_references_to_master(
    doc: Document,
    references: list[dict],
    normal_style: str,
) -> None:
    """
    Append real reference texts to the global REFERENCES chapter in the form:

      {REF:chapter03_local_001} Reference text...
    """
    if not references:
        return

    references_heading = find_heading_paragraph(doc, "REFERENCES")
    anchor = remove_placeholder_paragraphs_after_heading(references_heading)

    for ref in references:
        ref_id = ref.get("id", "REF:unknown")
        raw_text = (ref.get("raw_text") or "").strip()
        if not raw_text:
            continue

        line = "{" + ref_id + "} " + raw_text
        anchor = insert_paragraph_after(anchor, text=line, style_name=normal_style)

    page_break_para = insert_paragraph_after(anchor, style_name=normal_style)
    run = page_break_para.add_run()
    run.add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Table / equation insertion
# ---------------------------------------------------------------------------

def is_numbered_table_object_id(object_id: str) -> bool:
    """
    Return True for numbered table ids like 'chapter03_002' and False for
    unnumbered helper objects like 'chapter03_table_u002'.
    """
    return "_table_u" not in object_id


def build_table_cap_marker(table_block: dict) -> str:
    """
    Convert a numbered table block object_id like 'chapter03_001' to a caption marker
    '{TAB_CAP:chapter03_001}'.
    """
    object_id = table_block.get("object_id", "unknown")
    return "{TAB_CAP:" + object_id + "}"


def insert_table_at_anchor(
    anchor: Paragraph,
    table_block: dict,
    assets_doc_dir: Path,
    normal_style: str,
    table_style_name: Optional[str],
) -> Paragraph:
    """
    Insert a real table exactly at the current assembly anchor.

    Rules:
    - numbered tables get a {TAB_CAP:...} caption line
    - unnumbered table-like objects do not get a TAB_CAP field
    - if an unnumbered object has caption_text, render plain text only
    """
    object_id = table_block.get("object_id", "unknown")
    caption_text = (table_block.get("caption_text") or "").strip()
    table_file_rel = table_block.get("file", "")

    table_yaml_path = assets_doc_dir / table_file_rel
    rows = load_table_rows(table_yaml_path) if table_yaml_path.exists() else [["[Missing table asset]"]]

    caption_anchor = anchor

    if is_numbered_table_object_id(object_id):
        caption_line = build_table_cap_marker(table_block)
        if caption_text:
            caption_line += " " + caption_text
        else:
            caption_line += " "
        caption_anchor = insert_paragraph_after(
            anchor,
            text=caption_line,
            style_name=normal_style,
        )
    elif caption_text:
        caption_anchor = insert_paragraph_after(
            anchor,
            text=caption_text,
            style_name=normal_style,
        )

    table_obj, trailing_para = insert_table_after(caption_anchor, rows)

    if table_style_name:
        try:
            table_obj.style = table_style_name
        except Exception:
            pass

    return trailing_para


def insert_equation_placeholder_at_anchor(
    anchor: Paragraph,
    equation_block: dict,
    normal_style: str,
) -> Paragraph:
    """
    Insert a visible placeholder for an equation block.

    Step 3 behaviour:
    - keep equation position in the document
    - preserve equation numbering if available
    """
    equation_no = (equation_block.get("equation_no") or "").strip()

    placeholder = "[Equation"
    if equation_no:
        placeholder += f" {equation_no}"
    placeholder += " omitted]"

    return insert_paragraph_after(anchor, text=placeholder, style_name=normal_style)


# ---------------------------------------------------------------------------
# Assembly logic
# ---------------------------------------------------------------------------

def assemble_one_document(
    skeleton_path: Path,
    content_yaml_path: Path,
    assets_doc_dir: Path,
    output_path: Path,
) -> None:
    """
    Insert one normalized document into the skeleton and save the result.

    Current behaviour:
    - insert headings and paragraphs
    - preserve all inline fields exactly
    - insert real tables inline in normalized order
    - numbered tables get TAB_CAP fields
    - unnumbered table-like objects do not get TAB_CAP fields
    - equations are inserted as visible numbered placeholders
    - do NOT insert figures yet
    - do NOT insert a local References subsection inside the chapter
    - append bibliography entries to the global REFERENCES chapter
    """
    if not skeleton_path.exists():
        raise FileNotFoundError(f"Skeleton file not found: {skeleton_path}")

    if not content_yaml_path.exists():
        raise FileNotFoundError(f"Normalized content YAML not found: {content_yaml_path}")

    payload = load_yaml(content_yaml_path)

    title = payload.get("title")
    if not title:
        raise ValueError("Normalized content is missing 'title'.")

    blocks = payload.get("blocks", [])
    if not isinstance(blocks, list):
        raise ValueError("'blocks' must be a list in content.yaml")

    references = payload.get("references", [])
    if not isinstance(references, list):
        raise ValueError("'references' must be a list in content.yaml")

    doc = Document(skeleton_path)

    h1_style = find_first_existing_style(doc, ["H1", "Heading 1", "HEADING1"])
    h2_style = find_first_existing_style(doc, ["H2", "Heading 2", "HEADING2"], fallback=h1_style)
    h3_style = find_first_existing_style(doc, ["H3", "Heading 3", "HEADING3"], fallback=h2_style)
    normal_style = find_first_existing_style(doc, ["Normal", "normal"])
    table_style_name = find_first_existing_style(doc, ["Table Grid", "TableGrid", "Normal Table"], fallback="")
    if table_style_name == "":
        table_style_name = None

    # Insert main body content under the chapter heading.
    heading_para = find_heading_paragraph(doc, title)
    anchor = remove_placeholder_paragraphs_after_heading(heading_para)

    for block in blocks:
        block_type = block.get("type")

        if block_type == "heading":
            text = block.get("text", "")
            level = int(block.get("level", 1))

            if is_references_heading(text):
                continue

            # Skip generic source heading like "CHAPTER 3".
            if normalize_heading_text(text) == normalize_heading_text(title):
                continue

            # Replace the existing skeleton chapter heading text with the first
            # real level-1 heading from the source document.
            if level == 1:
                set_paragraph_text(heading_para, text)
                continue

            anchor = insert_heading_block(
                anchor=anchor,
                text=text,
                level=level,
                h1_style=h1_style,
                h2_style=h2_style,
                h3_style=h3_style,
            )
            continue

        if block_type == "paragraph":
            anchor = insert_paragraph_block(
                anchor=anchor,
                text=block.get("text", ""),
                style_name=normal_style,
            )
            continue

        if block_type == "table":
            anchor = insert_table_at_anchor(
                anchor=anchor,
                table_block=block,
                assets_doc_dir=assets_doc_dir,
                normal_style=normal_style,
                table_style_name=table_style_name,
            )
            continue

        if block_type == "equation":
            anchor = insert_equation_placeholder_at_anchor(
                anchor=anchor,
                equation_block=block,
                normal_style=normal_style,
            )
            continue

        if block_type == "figure":
            # Ignore figures for now.
            continue

        anchor = insert_paragraph_block(
            anchor=anchor,
            text=f"[Unsupported block type: {block_type}]",
            style_name=normal_style,
        )

    insert_page_break_after(anchor, style_name=normal_style)

    append_references_to_master(
        doc=doc,
        references=references,
        normal_style=normal_style,
    )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------

def parse_args() -> argparse.Namespace:
    """
    Parse command-line arguments.
    """
    parser = argparse.ArgumentParser(
        description="Insert one normalized NACIE TECDOC document into the master skeleton."
    )
    parser.add_argument(
        "content_yaml",
        help="Path to content.yaml produced by normalize_docx.py",
    )
    parser.add_argument(
        "--skeleton",
        default=str(DEFAULT_SKELETON_PATH),
        help="Path to the skeleton master .docx",
    )
    parser.add_argument(
        "--output",
        default=str(DEFAULT_OUTPUT_PATH),
        help="Path to output assembled .docx",
    )
    return parser.parse_args()


def main() -> int:
    """
    Run the insertion prototype.
    """
    args = parse_args()

    content_yaml_path = Path(args.content_yaml).resolve()
    skeleton_path = Path(args.skeleton).resolve()
    output_path = Path(args.output).resolve()

    document_dir = content_yaml_path.parent
    section_dir = document_dir.parent
    section_name = section_dir.name
    document_id = document_dir.name
    assets_doc_dir = ROOT / "04-assets" / section_name / document_id

    assemble_one_document(
        skeleton_path=skeleton_path,
        content_yaml_path=content_yaml_path,
        assets_doc_dir=assets_doc_dir,
        output_path=output_path,
    )

    print(f"Skeleton input : {skeleton_path}")
    print(f"Content input  : {content_yaml_path}")
    print(f"Output written : {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())